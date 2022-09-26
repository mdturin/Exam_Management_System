import datetime
import io

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm

from account.models import *
from dashboard.models import *

doc = Document()

section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE

section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Mm(297)
section.page_height = Mm(210)

title = doc.add_heading('Duty Roaster List', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows=1, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

row = table.rows[0].cells
row[0].add_paragraph("Date").alignment = WD_ALIGN_PARAGRAPH.CENTER
row[1].add_paragraph("Course Code").alignment = WD_ALIGN_PARAGRAPH.CENTER
row[2].add_paragraph("Supervisor").alignment = WD_ALIGN_PARAGRAPH.CENTER
row[3].add_paragraph("Examiners").alignment = WD_ALIGN_PARAGRAPH.CENTER
# row[4].add_paragraph("Assistance").alignment = WD_ALIGN_PARAGRAPH.CENTER


def get_data(exams: Exam):
    exams_data = []
    for exam in exams:
        data = []
        data.append(str(exam.exam_date))
        data.append(exam.course.code)
        data.append(exam.supervisor.get_name)
        data.append(list(exam.examiners.all()))
        exams_data.append(data)
    return exams_data


def roaster(exams: Exam, routine: Routine):
    data = get_data(exams)
    for date, code, supervisor, examiners in data:
        row = table.add_row().cells
        row[0].add_paragraph(date).alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].add_paragraph(code).alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].add_paragraph(supervisor).alignment = WD_ALIGN_PARAGRAPH.CENTER
        # row[3].add_paragraph("\n".join(examiners)
        #                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        # row[4].add_paragraph("\n".join(assistance)
        #                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f"Sample/Roaster/{routine.name}.docx")
